from decouple import config
import subprocess
import docx
import re
import os
from os.path import join
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pypdf import PdfReader, PdfWriter


def docx_format_string(string):
    """
    Convert `string` to an appropriately formatted paragraph.
    :param string: This argument must be a str or a NoneType, otherwise an exception is raised. If the string is a
    NoneType then just return it.
    :return str: This is the formatted argument. If the argument was None then None should be returned.
    """
    # Check if string is None or not a string
    if string is not None and not isinstance(string, str):
        raise ValueError("The argument to `docx_format_string` must be a str or None")

    # Return the argument if it is a NoneType.
    if string is None:
        return ""

    # Convert CRLF text to LF
    string = re.sub(r"\r\n|\n", r"\n", string)

    # Remove leading and trailing whitespaces
    string = string.strip()
    #string = unidecode_expect_ascii(string, errors="preserve", exclude_chars=["•", b'\xe2\x9c\x94'.decode('utf-8'), '“', '”'])

    return string


def docx_substr_in(substring, main_string):
    """
    Indicate whether a specified substring is in another string. If substring is null then return False (this is the
    meaningful distinction between this function and the regular python built-in "in" membership operator.
    :param substring: This is a str that we're searching an occurrence of inside the main string.
    :param main_string: This is a str within which we are searching for substring.
    :return: Return a boolean indicating that either the substring was found inside `main_string`.
    """
    # Check if substring is None or not a string
    if substring is None or not isinstance(substring, str):
        # If substring is None or not a string, return False
        return False

    # Check if main_string is None or not a string
    if main_string is None or not isinstance(main_string, str):
        # If main_string is None or not a string, return False
        return False

    # Check if substring is an empty string
    if not substring:
        # If substring is an empty string, return False
        return False

    # Use the "in" operator to check substring in main_string
    return substring in main_string


def docx_format_number(number, val_to_return_if_null=None, return_str=False):
    """
    Convert `number` to a string and formatted as so:
    If `number` is an int then simply convert it If `number` is a float but also a whole number then remove the decimal
    If `number` is a float with more than 6 digits after the decimal then slice digits off past the 6
    If `number` is a str then check if it's the empty string.
    If so assign it to val_to_return_if_null and return it. Otherwise check if `number` is numeric. If so, convert it to either a
    float or an int. Otherwise, raise a ValueError.

    :param number: - This must be an int or float, otherwise an exception will be raised.
    :param val_to_return_if_null: - This value is returned if `number` is None or the empty string.
    :param return_str: - This is a boolean value indicating whether to the returned value should be a str or not. If
    the passed-in value is a string thi
    :return: str- This is a formatted str representation of the argument.
    """
    if not isinstance(return_str, bool):
        raise Exception("The `return_str` argument must be a bool.")

    if not isinstance(number, (int, float, str, type(None))):
        raise ValueError("The `number` argument must be an int, float, str, or NoneType.")

    if isinstance(number, type(None)):
        return val_to_return_if_null if not return_str else str(val_to_return_if_null)

    if isinstance(number, str):
        if number == "":
            return val_to_return_if_null if not return_str else str(val_to_return_if_null)
        try:
            number = int(number)
        except ValueError:
            try:
                number = float(number)
            except ValueError:
                if isinstance(val_to_return_if_null, type(None)):
                    raise ValueError(
                        "A str p_type passed to `docx_format_number` must be numeric otherwise a non-NoneType argument "
                        "for `val_to_return_if_null` must be supplied")
                else:
                    return val_to_return_if_null

    if isinstance(number, float):
        if number.is_integer():
            if return_str:
                return str(int(number))
            else:
                return int(number)
        if return_str:
            return "{:.6f}".format(number).rstrip("0").rstrip(".")
        else:
            return float("{:.6f}".format(number))

    elif isinstance(number, int):
        if return_str:
            return str(number)
        else:
            return number


def docx_format_seismic_reflection_data(location, position, position_type, description):
    """
    Return a string containing appropriately formatted seismic reflection data given the four relevant data fields.
    :location str: This string should be a value from a data field.
    :position str: This string should be a value from a data field.
    :position_type str: This string should be a value from a data field.
    :description str: This string should be a value from a data field.
    :return str: This is the formatted string composed of data from the above fields.
    """
    if not all(isinstance(arg, (str, type(None))) for arg in (location, position, position_type, description)):
        raise TypeError("All arguments to `docx_format_seismic_reflection_data` must be strings.")

    location = location if location is not None else ""
    position = position if position is not None else ""
    position_type = position_type if position_type is not None else ""
    description = description if description is not None else ""

    return_me = ""  # Progressively append contents to this str and finally return it.
    if location:
        return_me += "Line: " + location
    if position:
        return_me += "Position: " if return_me == "" else "\n" + "Position"
        if position_type:
            return_me += " " + position_type.upper()
        return_me += " " + position
    if description:
        return_me += description if return_me == "" else "\n" + description
    return return_me


def docx_convert_word_to_pdf(input_path, output_path):
    """
    Run the PDF converter tool on the input word document and write a converted PDF file to the out_path arg.
    :param input_path: This str is an absolute path to the Word document to be converted.
    :param output_path: This str is an absolute path to the desired location to place the converted PDF.
    :return: None
    """
    # Read the location of the LibreOffice python executable from the .env file.
    path = config("LIBRE_OFFICE_PYTHON_INSTALLATION_PATH")
    cmd = [
        path,
        "-c",
        "import unoserver.client as t; t.converter_main()",
        input_path,
        output_path
    ]

    # Run the command as a shell script.
    print("Running the following command to convert completed webform templates to PDF:")
    print(" ".join(cmd))
    subprocess.run(cmd)


def docx_delete_paragraph(paragraph):
    """
    Delete a paragraph object in a Document object.
    :param paragraph: This is a Paragraph object to be deleted.
    :return: None
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def docx_delete_paragraphs(doc_obj, indices):
    """
    Delete a subset of Paragraph objects from the specified Document (doc_obj) specified by `indices`
    :param doc_obj: This is the Document object from which paragraphs will be deleted.
    :param indices: This is a list of indices indicating which paragraph item in the ordering defined by the
    `iter_block_items` generator function to delete
    :return: None
    """
    indices_sorted = indices
    indices_sorted.sort(reverse=True)
    document_block_items = [element for element in iter_block_items(doc_obj)]
    for index in indices_sorted:
        if isinstance(document_block_items[index], docx.text.paragraph.Paragraph):
            docx_delete_paragraph(document_block_items[index])
        elif isinstance(document_block_items[index], docx.table.Table):
            docx_delete_table(document_block_items[index])


def docx_delete_table(table):
    """
    Delete a Table object from a Document object.
    :param table: This is a table object to be deleted.
    :return: None
    """
    table._element.getparent().remove(table._element)


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order. Each returned value is an instance of
    either Table or Paragraph. *parent* would most commonly be a reference to a main Document object, but also works for
    a _Cell object, which itself can contain paragraphs and tables.
    :param parent: This is either a Document object or a Cell object. If neither of these are the case then an exception
    is raised.
    :yield: This function yields every child element of the parent in document-order. I.E. in the order in which they
    occur in the document.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def docx_redact_conditional(doc_obj, search_data):
    """
    Search the document for a conditional field and either delete the paragraphs between the conditional or just
    remove the conditional tags. Conditional tags are of the form `{% conditional-section xxx %}` where `xxx` matches
    a key in `search_data`. The template tag is delimited via the tag: `{% end-section %}`. If the corresponding
    value in `search_data` is False then delete both template tags along with all paragraph"s and tables that appear
    between them. Otherwise, just delete the template tag and leave everything between them alone.
    :param doc_obj: This is the Document object that contains the tags.
    :param search_data: This is a dictionary of template tag names mapped to Boolean values.
    :return: None
    """
    if not isinstance(doc_obj, docx.document.Document):
        raise RuntimeError("The `doc_obj` argument must be a Document-p_type")
    if not isinstance(search_data, dict):
        raise RuntimeError("The `search_data` argument must be a dictionary")

    paragraph_indices_to_delete = []  # Rather than delete paragraphs as we"re reading them store the indices later and
    # delete them
    opening_tag = re.compile(r"{%\s*conditional-section\s+(.*?)\s*%}")
    closing_tag = re.compile(r"{%\s*end-section\s*%}")
    document_block_items = [element for element in iter_block_items(doc_obj)]  # Store a reference to all items in the
    # order they appear in the document
    for index, p in enumerate(document_block_items):
        if not isinstance(p, docx.text.paragraph.Paragraph):  # Search paragraphs for conditional tags. Tags do not
            # appear in tables so skip them
            continue
        run_sum = ""
        for i in range(len(p.runs)):
            run_sum += p.runs[i].text
        if len(p.runs) > 0:
            match_open_tag = opening_tag.search(run_sum)
            match_close_tag = closing_tag.search(run_sum)
            if match_open_tag and match_close_tag:  # If an opening tag is discovered and
                if search_data.get(re.sub(r"(.*?)\{%\s*conditional-section\s*|\s*%}(.*?)", "",
                                          match_open_tag.group())):  # If opening and closing tags are both present
                    # and the tag evaluates to True in the database then simply delete the template tags
                    p.runs[0].text = re.sub(r"{%\s*(conditional-section\s+\S+|end-section)\s*%}", "", run_sum)
                    for j in range(len(p.runs) - 1, 0,
                                   -1):  # This idiom will be repeated a lot and just means "delete any text runs
                        # after the first".
                        p._p.remove(p.runs[j]._r)
                else:
                    if run_sum[:match_open_tag.start()] + run_sum[
                                                          match_close_tag.end():] == "":  # If this paragraph
                        # consists solely of the opening, closing tags and text in between then just delete the
                        # entire paragraph element.
                        paragraph_indices_to_delete.append(index)
                    else:
                        p.runs[0].text = run_sum[:match_open_tag.start()] + run_sum[
                                                                            match_close_tag.end():]  # Remove both
                        # tags and all text in between both tags.
                        for j in range(len(p.runs) - 1, 0, -1):
                            p._p.remove(p.runs[j]._r)

            elif match_open_tag and not match_close_tag:  # If we've found an opening conditional but not a closing
                # conditional tag...
                if search_data.get(re.sub(r"(.*?)\{%\s*conditional-section\s*|\s*%}(.*?)", "",
                                          match_open_tag.group())):  # If closing tag is not found and the boolean is
                    # True then just delete the tag and search for the closing tag to delete
                    if re.sub(r"{%\s*conditional-section\s+(.*?)\s*%}", "",
                              run_sum) == "":  # If this paragraph contains just the template tag then delete the
                        # paragraph entirely
                        paragraph_indices_to_delete.append(index)
                    else:
                        p.runs[0].text = re.sub(r"{%\s*conditional-section\s+(.*?)\s*%}", "",
                                                run_sum)  # Otherwise just delete the template tag and leave
                        # everything else
                    for future_p in range(index + 1,
                                          len(document_block_items)):  # Fetch the closing tag in future runs and
                        # delete it
                        if not isinstance(document_block_items[future_p],
                                          docx.text.paragraph.Paragraph):  # Skip over tables when looking for the
                            # future closing tag
                            continue
                        inner_sum = ""
                        for i in range(len(document_block_items[
                                               future_p].runs)):  # Combine all runs in this paragraph together so we
                            # can parse the paragraph easily
                            inner_sum += document_block_items[future_p].runs[i].text
                        if closing_tag.search(inner_sum):
                            if re.sub(closing_tag, "", inner_sum) == "":
                                paragraph_indices_to_delete.append(future_p)
                            else:
                                document_block_items[future_p].runs[0].text = re.sub(closing_tag, "", inner_sum)
                            break

                else:  # Otherwise we want to delete not just the tags but everything in between as well
                    if run_sum[:match_open_tag.start()] == "":
                        paragraph_indices_to_delete.append(index)
                    else:
                        p.runs[0].text = run_sum[:match_open_tag.start()]  # Delete the tag and all characters after it
                    # Find the future closing tag and keep deleting character until that tag is reached and then
                    # delete it too
                    future_p = index + 1
                    while future_p < len(document_block_items):
                        if not isinstance(document_block_items[future_p],
                                          docx.text.paragraph.Paragraph):  # Skip over tables when looking for the
                            # future closing tag
                            paragraph_indices_to_delete.append(
                                future_p)  # We"ve encountered a table which means we haven"t found our closing tag
                            # so add the table to the list of stuff to delete
                            future_p += 1
                            continue  # And skip over the table in our search for the closing tag
                        inner_sum = ""
                        for i in range(len(document_block_items[future_p].runs)):
                            inner_sum += document_block_items[future_p].runs[i].text
                        if closing_tag.search(inner_sum):
                            if inner_sum[closing_tag.search(inner_sum).end():] == "":
                                paragraph_indices_to_delete.append(future_p)
                            else:
                                document_block_items[future_p].runs[0].text = inner_sum[closing_tag.search(
                                    inner_sum).end():]  # Delete all characters appearing before the closing tag
                            break
                        else:
                            paragraph_indices_to_delete.append(future_p)
                            future_p += 1
                            # docx_delete_paragraph(doc_obj.paragraphs[future_p]) doc_obj.paragraphs[future_p].runs[
                            # 0].text = ""  # If the closing tag isn't in this paragraph then just delete the entire
                            # paragraph
    # Delete all paragraphs by specifying the indices that were previously marked for deletion and invoking this helper.
    docx_delete_paragraphs(doc_obj, paragraph_indices_to_delete)


def docx_build_table(doc_obj, table_id, table_data, background_color="#FFFFFF"):
    """
    Replace all matching `regex` of the table specified by `id` with `replace`
    :param doc_obj: This is the document object to modify the specified table of
    :param table_id: The id of the table corresponds to the order that it appears in the docx (0-indexed)
    :param table_data: This is the data that we will populate the table indicated by `table_id` with. This is an array
        of tuples.
    :param background_color: This str should be a hexadecimal value (one starting with "#" and 6 digits) specifying the
        color to paint the cells of the new table.
    :return:
    """
    # Get the table on the Word document that we will be building out
    table = doc_obj.tables[table_id]
    # Create a new row in the existing Word document table for every row in the data that is to be added
    for incoming_row in table_data:
        row = table.add_row()
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement("w:cantSplit"))
        new_row = row.cells  # Create a new row to be added to the word table
        for i in range(len(incoming_row)):  # For every column of data to be added...
            new_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = new_row[i].paragraphs[0]  # Delete the paragraph that's auto-inserted when a cell is created
            docx_delete_paragraph(paragraph)
            new_row[i].add_paragraph(str(incoming_row[i]), doc_obj.styles["Proponent Table Paragraph"])
            # Apply shading to the table cell.
            shading_elm = parse_xml((r'<w:shd {} w:fill="' + background_color + r'"/>').format(nsdecls("w")))
            new_row[i]._tc.get_or_add_tcPr().append(shading_elm)


def docx_search_and_replace_tags(doc_obj, search_data, external_i=None, crash_on_except=False):
    """
    Search the specified document for tags and replace these tags with a corresponding value. The tags and values
    are stored in a dict corresponding to the `search_data` parameter. Dict keys correspond to tag values expected to
    be present in the document to be search (specified by `doc_obj`). Dict values are either strings or lists of
    strings. They can also be 2-tuples where the first value is exactly what was specified in the previous sentence
    and the second value is a dictionary of strings mapping to other strings. The mapping serves as a translation table.
    If the first value of the 2-tuple matches any keys in the second value then the value associated with that key as
    specified in the 2-tuple's second value dictionary is returned. I.E. If this tuple is passed in the output is
    '0' because the first 2-tuple value evaluated to 'a' and 'a' was translated to '0': (['a', 'b', 'c'][0], {'a': '0', 'b': '1'})
    If a value is a list of strings then `external_i` must be specified in order to index into the value and
    retrieve a specific string. If there are multiple values that are list-types then the same external_i value will be
    used to index into each of them. This function should be called once per document-template-tag-dict pair. The
    recursion in this function is used to search through paragraphs nested inside of table cells, headers, and footer
    elements.
    :param doc_obj: This is the Document object to search and replace tags through.
    :param search_data: This is a dictionary of template tag values paired with either strings or a list of strings.
    :param external_i: This is an int used to index into the list of strings.
    :return: None
    """
    def sub_cb(match):
        match_parsed = re.sub(r"\{\{|\}\}", "", match.group())
        val = search_data.get(match_parsed, None)
        val_bindings = {}
        if isinstance(val, tuple):
            temp = val[1]
            val = val[0]
            val_bindings = temp
        if val is not None:
            if isinstance(val, list):
                if not isinstance(external_i, int):
                    raise ValueError("A list value was found in the tag dictionary but either no external_i arg was"
                                     "specified or the specified external_i arg was not an int.")
                if external_i >= len(val):
                    if crash_on_except:
                        raise IndexError(
                            """The `external_i` arg tried to index outside of a list value in the `search_data` arg."""
                        )
                    val = ""
                else:
                    val = str(val[external_i])
            else:
                val = str(val)
        if val_bindings:    # This "if" is necessary because user might not have specified a translation dict causing the line below to crash
            val = val_bindings.get(val) if val in val_bindings else val
        return val

    # Verify argument types.
    if not isinstance(doc_obj, (docx.document.Document, docx.section._Header, docx.section._Footer, docx.table._Cell)):
        raise ValueError("The doc_obj argument must be a Document.")
    if not isinstance(search_data, dict):
        raise ValueError("The search_data argument must be a dictionary.")
    if not isinstance(external_i, (type(None), int)):
        raise ValueError("The external_i argument must be an int.")

    # Search all text in the document for regex matches except for header content and tables
    tags = re.compile(
        "|".join(map(re.escape, ["{{" + key + "}}" for key in search_data])))  # Define set of re patterns to match on
    for p in doc_obj.paragraphs:
        run_sum = ""
        for i in range(len(p.runs)):
            run_sum += p.runs[i].text
        if len(p.runs) > 0:
            if run_sum != "":
                new_text = docx_format_string(tags.sub(sub_cb, run_sum))
                p.runs[0].text = new_text
                if new_text == "☑" or new_text == "☐":
                    p.runs[0].font.size = Pt(16)
                    p.runs[0].font.bold = True
                # p.runs[0].font.highlight_color = RGBColor(238, 238, 238)
            for j in range(len(p.runs) - 1, 0,
                           -1):  # Delete previous runs (in reverse order) now that we have a full match
                p._p.remove(p.runs[j]._r)

    # Search all headers and footers for matching text and replace
    if isinstance(doc_obj, docx.document.Document):
        for doc_header in doc_obj.sections:
            docx_search_and_replace_tags(doc_header.header, search_data, external_i)
            docx_search_and_replace_tags(doc_header.footer, search_data, external_i)

    # Search all tables for matching text and replace
    if isinstance(doc_obj, docx.document.Document):
        for doc_table in doc_obj.tables:
            for table_row in doc_table.rows:
                for table_cell in table_row.cells:
                    docx_search_and_replace_tags(table_cell, search_data, external_i)


def docx_define_styles(doc_obj):
    """
    Define new styles to be applied to content inside a Word document. This function does not apply any of the styles
    to the Document arg but rather defines them as "selectable" styles within the Document to be applied later. As of
    12/28/2023 This function only defines the "Proponent Table Paragraph" style which is used when constructing tables
    via the docx_build_table function. In general this particular style can be accessed with the following:
    doc_obj.styles["Proponent Table Paragraph"]
    :param doc_obj: The Word document object to define the new styles for
    :return:
    """
    if not isinstance(doc_obj, docx.document.Document):
        raise ValueError("The doc_obj argument must be a Document.")

    # Define the "Proponent Table" style for the Document arg. This style should always be the one used on tables
    # (not just the proponent table) unless there is a good reason not to do so.
    proponent_table_paragraph_style = doc_obj.styles.add_style("Proponent Table Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    proponent_table_paragraph_style.font.name = "Arial"
    proponent_table_paragraph_style.font.bold = False
    proponent_table_paragraph_style.font.italic = False
    proponent_table_paragraph_style.font.size = Pt(9)
    proponent_table_paragraph_style.font.shadow = False
    proponent_table_paragraph_style.font.color.rgb = RGBColor(51, 25, 0)
    proponent_table_paragraph_style.paragraph_format.space_before = Pt(0)


def docx_append_pages(reader, writer, title, idx, parent_outline=None, create_parent_outline_name=None):
    """
    Merge all pages in the reader object into the writer object at a specified location and create a bookmark
    for the first page of the reader object.
    :param file_dir: This is a str that's the path to the directory where a new version of the file specified for the
    reader param only in the event that the reader param file is corrupt and needs to be repaired.
    :param file_name: This is the name of the file that was specified as the argument to `reader` when `reader` was
    instantiated. This is name is necessary because it specifies the name of the file to be repaired if
    :param reader: This is the user upload.
    :param writer: This is a writer object the user upload will be appended to.
    :param title: This is a str indicating what to name the bookmark.
    :param idx: This is an integer indicating which index location in the writer to append the new user upload
    :param parent_outline: This is an object that's used as an argument for `add_outline_item` in order to
    specify a parent bookmark for the new bookmark.
    :param create_parent_outline_name: This is a str indicating the name for a parent element to be created
    and assigned the parent of the new bookmark (which is specified by `title` param). This arg cannot be set
    if param `parent_outline` is also set. Doing so will raise an exception.
    pages to
    :return: A 2-tuple with the first element being an int equal to the number of pages in the user upload and
    the second element being the parent_outline object that was either specified as a param or created via the
    `create_element_outline_name` parameter.
    """

    if isinstance(reader, type(None)):
        return 0, parent_outline

    if parent_outline and create_parent_outline_name:
        raise RuntimeError("`docx_bookmark_and_process_pages` function call args `parent_outline` and "
                           "`create_parent_outline` cannot both be set. Set either one or the other or leave "
                           "both as None")

    writer.append(reader, import_outline=False)

    if create_parent_outline_name:
        parent_outline = writer.add_outline_item(create_parent_outline_name, idx)
    writer.add_outline_item(title, idx, parent_outline)
    return len(reader.pages), parent_outline


def docx_bookmark_and_process_pages(reader, writer, title, idx, offset, ident_target, ident_empty,
                                    ident_target_cont=None, parent_outline=None, create_parent_outline_name=None):
    """
    Write pages from the PdfReader file to the PdfWriter file and bookmark it.
    :param reader: This is a PdfReader representing the source PDF from whose pages we will write to the writer file.
    :param writer: This is a PdfWriter representing a PDF that the reader pages will be written to.
    :param title: This is str indicating the title for the PDF outline (bookmark) that will be created for the incoming
    pages from the reader param.
    :param idx: This is an int denoting the page in the reader param to start reading from.
    :param offset: This is an int offset for bookmarks (necessary because user uploads aren't factored). The offset is
    necessary for determining which page of the writer should be bookmarked.
    :param ident_target: This is a string used to identify the coversheet page. This string must be a substring
    of the result of calling reader.pages.extract_text() on the coversheet page.
    :param ident_target_cont: This is a str used to identify pages that come after the `ident_target`. If this
    substring isn't specified then no future page will be searched. Otherwise, execution will repeat until
    every repeat of this "continued" page is located. All repeats of this page in the `reader` must have this
    str as a substring.
    :param ident_empty: This is a string used to identify pages that are empty. Pages that are identified as
    empty will not be added to the writer.
    :param parent_outline: This is an object that's used as an argument for `add_outline_item` in order to
    specify a parent bookmark for the new bookmark.
    :param create_parent_outline_name: This is a str indicating the name for a parent element to be created
    and assigned the parent of the new bookmark (which is specified by `title` param). This arg cannot be set
    if param `parent_outline` is also set. Doing so will raise an exception.
    :return: A 3-tuple is returned where the first element indicates the next page-index in `reader` to process,
    the second element is an int indicating the offset that should be added to the first element for
    bookmark placement, and the third is the parent_outline_item that was either passed in or was created
    by specifying the `create_parent_outline_item` param.
    """
    target_found = False  # Indicate if the target page has been found and bookmarked
    target_cont_found = False
    ret_idx = idx  # Copy these two ints to modify them and return them later.
    ret_offset = offset

    if isinstance(reader, type(None)):  # Do nothing if the reader is None (likely due to its file not being
        # found
        return ret_idx, ret_offset, parent_outline

    if parent_outline and create_parent_outline_name:
        raise RuntimeError("`docx_bookmark_and_process_pages` function call args `parent_outline` and "
                           "`create_parent_outline` cannot both be set. Set either one or the other or leave "
                           "both as None")

    for i in range(idx, len(reader.pages)):
        page_text = reader.pages[i].extract_text()
        if i + 1 < len(reader.pages):
            next_page_text = reader.pages[i + 1].extract_text()
        else:
            next_page_text = ""

        if page_text == ident_empty:  # If the page text matches what is expected for an empty page then
            # skip it.
            ret_idx += 1  # Advancing the return index is necessary to keep future calls of this function
            # up-to-date with where to start processing pages in the reader.
            ret_offset -= 1  # This variable is a cursor for the writer object and the writer object knows
            # nothing about the blank page so the ret_idx advance has to be negated for the bookmark location
            # to be correct.
            continue

        if next_page_text == ident_empty:  # If the next page is empty then look ahead until it's not empty.
            # If there are no non-empty pages left in the document assign next_page_text to the empty string.
            for j in range(i + 1, len(reader.pages)):
                try:
                    if reader.pages[j].extract_text() != ident_empty:
                        next_page_text = reader.pages[j].extract_text()
                        break
                except IndexError:
                    next_page_text = ""
                    break

        writer.add_page(reader.pages[i])  # Now we know this isn't a blank page so add it to the final document.

        # This is the expected situation when we've located the target and there is no "continued" page
        # specified
        if ident_target in page_text and not docx_substr_in(ident_target_cont, next_page_text):
            if create_parent_outline_name:
                parent_outline = writer.add_outline_item(create_parent_outline_name, ret_idx + ret_offset)
            writer.add_outline_item(title, ret_idx + ret_offset, parent_outline)  # Bookmark the page and return.
            return ret_idx + 1, ret_offset, parent_outline

        # This is the expected situation for when we've found the target and there is a defined "continued"
        # page.
        elif ident_target in page_text and docx_substr_in(ident_target_cont, next_page_text):
            if create_parent_outline_name:
                parent_outline = writer.add_outline_item(create_parent_outline_name, ret_idx + ret_offset)
            writer.add_outline_item(title, ret_idx + ret_offset, parent_outline)
            target_found = True
            target_cont_found = True

        # This case is rarer because it involves there being multiple continued pages and looking in the middle
        # of one of these sets of pages. Otherwise, if this case is true then something is awry!
        elif ident_target not in page_text and docx_substr_in(ident_target_cont, next_page_text):
            if target_cont_found and target_found:
                pass
            else:
                raise RuntimeError("`docx_bookmark_page` ERROR: A \"cont\" page was found before the target.")

        # This is the case where none of the specified pages were found on this scan, so hopefully they already
        # have been located otherwise something has gone wrong.
        elif ident_target not in page_text and not docx_substr_in(ident_target_cont, next_page_text):
            if target_found and target_cont_found:
                return ret_idx + 1, ret_offset, parent_outline
            else:
                pass  # We haven't found any targets yet so keep looking.

        ret_idx += 1

    raise RuntimeError("`docx_bookmark_page` ERROR: EOF for the PdfReader was reached and the targets "
                       "weren't all found.")


def get_safely(el, i, return_empty_string_on_null=False):
    """
    Safely access the i'th element of el without worry of crashing.
    :param el: This is a list or tuple that the caller wants to index into.
    :param i: This is an int serving as the index.
    :param return_empty_string_on_null: This flag signals the program to return "" instead of None if either the
    list/tuple access fails or if the access element is None.
    :return:
    """
    if (isinstance(el, list) or isinstance(el, tuple)) and len(el) > 0:
        if 0 <= i < len(el):
            if return_empty_string_on_null:
                if isinstance(el[i], type(None)):
                    return ""
            return el[i]
    if return_empty_string_on_null:
        return ""
    else:
        return None


def validate_path(path, crash_on_miss=False):
    """
    Validates the existence of a directory at the specified path. If the directory
    does not exist, it attempts to create it and prints a warning message.

    Parameters:
    - path (str): The path of the directory to validate/create. Any other variable
    p_type will result in a runtime error.

    Returns:
    - str: This is the exact same string that was passed in
    """
    if not os.path.exists(path):
        if crash_on_miss:
            raise RuntimeError("The specified path \"" + path + "\" was not found on the system.")
        if not isinstance(path, str):
            raise RuntimeError("An argument of p_type other than str was passed in.")
        print("WARNING: The following path doesn't exist: " + path)
        print("I will attempt to create this path...", end="")
        os.makedirs(path)
        print("\tSuccessful")
    return path


def instantiate_pdf_reader(pdf_dir, pdf_name, label=None):
    """
    Create a python PDF object.
    :param pdf_dir: This is a str path to the directory containing the PDF
    :param pdf_name: This is the filename of the PDF.
    :param except_on_missing: This is a boolean that causes an exception to be raised if the pdf wasn't found if True.
    :param label: This is an optional label used to identify the file in an error message if it's missing.
    :return: This function returns a PdfReader which is just a python representation of a PDF. If the path arguments
    are invalid and a pdf couldn't be found to instantiate then this function returns None.
    """
    try:
        # Make sure the supplied filename and path are not None
        if not pdf_dir or not pdf_name:
            raise FileNotFoundError

        reader = PdfReader(join(pdf_dir, pdf_name))
        writer = PdfWriter()
        writer.append(reader)

    except FileNotFoundError:
        if label:
            print("ERROR: input file %s couldn't be located." % label)
        else:
            print("ERROR: input file %s couldn't be located." % pdf_name)
        return None
    except Exception:
        print("""ERROR: The input pdf is malformed.""")
        print("Attempting to repair the PDF...")
        proc = subprocess.run(["gs",
                               "-o",
                               join(pdf_dir, "TEMP_" + pdf_name),
                               "-sDEVICE=pdfwrite",
                               "-dPDFSETTINGS=/prepress",
                               join(pdf_dir, pdf_name)
                               ])
        # If the PDF repair attempt is successful then retry appending the page.
        if proc.returncode == 0:
            print("\nAttempt to repair: Successful")
            reader = PdfReader(join(pdf_dir, "TEMP_" + pdf_name))
        else:
            print("\nAttempt to repair: Unsuccessful")
            print("Treating file " + pdf_name + " as missing.")
            return None     # If the file cannot be repaired then treat it as if it was missing.

    return reader