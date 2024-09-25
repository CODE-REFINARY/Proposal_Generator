from pdf_gen_helper_functions import *
from os.path import join
from psycopg2 import connect
from psycopg2 import sql
from docx import Document
from datetime import datetime
import glob

class WordProposalGenerator:
    """
    This class is used to generate Word documents that will be later converted to PDFs. This class is capable of
    reading from the database, reading Word template files, and populating those templates with values based on the
    defined templating language. The templating language works as follows:
    Instances of {{ variable }} in the Word template will be searched for and matched with any defined tags that
    match it. If there is a match then this variable will be replaced. This is done using the
    docx_search_and_replace_tags function in the Word static library. This class reads data from a .env file which
    must be in the same directory as this file.
    """
    def __init__(self, pid):
        self.PROPOSAL_ID = pid
        self.PDF_GEN_DIR = validate_path(config("PDF_GEN_DIR"))
        self.LIBRE_OFFICE_PYTHON_INSTALLATION_PATH = validate_path(config("LIBRE_OFFICE_PYTHON_INSTALLATION_PATH"), True)
        self.PROPOSALS_BASE_DIR = validate_path(config("PROPOSALS_BASE_DIR"))
        self.PROPOSAL_DIR = validate_path(join(self.PROPOSALS_BASE_DIR, self.PROPOSAL_ID))
        self.PDF_UPLOADS_DIR = validate_path(join(self.PROPOSAL_DIR, "pdf_uploads"))
        self.conn = connect(  # This variable links the iodpdatadev database to this program
            host=config("DB_HOST"),
            database=config("DB_DATABASE"),
            user=config("DB_USERNAME"),
            password=config("DB_PASSWORD"),
            port=config("DB_PORT")
        )
        self.cur = self.conn.cursor()
        # This list keeps track of queries that have already been performed so that we don't needlessly access the
        # database with redundant queries.
        self.tables_queried = []

        # This list keeps track of all Word templates for sites that have been generated. It is a list of strings and
        # the strings are filenames.
        self.site_file_names = []

        # This text is will appear at the bottom of the document.
        self.footer_text = 'Generated: ' + datetime.now().isoformat(timespec='milliseconds')


    def get_db_rows(self, table):
        """
        Access the database and get specific rows from it based on a pre-defined query that's hardcoded into this function.
        This query is specified by supplying an argument to the table param.
        :param table: This is a str identifier that select what table and what query to run. It's matched againt a list
        of values in a switch-like statement and the result is that a specific query is run.
        :return: This function returns a 2-tuple where the first element is the result of the query and the second element
        is a list of all columns (in order) that were in the table specified
        """
        # If this table has already been queried then exit silently because it's member variable will already exist
        if table in self.tables_queried:
            return

        if table == "COVERSHEET":
            self.cur.execute("SELECT * FROM coversheet WHERE proposal_id = %s", (self.PROPOSAL_ID,))
            self.COVERSHEET_ROW = self.cur.fetchone()
            self.COVERSHEET_COLS = [desc[0] for desc in self.cur.description]

        elif table == "PROPOSAL":
            self.cur.execute("SELECT * FROM proposal WHERE id = %s", (self.PROPOSAL_ID,))
            self.PROPOSAL_ROW = self.cur.fetchone()
            self.PROPOSAL_COLS = [desc[0] for desc in self.cur.description]

        elif table == "MAIN_TEXT_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'DOC'", (self.PROPOSAL_ID,))
            self.MAIN_TEXT_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "SAFETY_REVIEW_REPORT_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'SRR'", (self.PROPOSAL_ID,))
            self.SAFETY_REVIEW_REPORT_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "CURRICULUM_VITAE_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'CV'", (self.PROPOSAL_ID,))
            self.CURRICULUM_VITAE_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "REVIEWERS_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'PR'", (self.PROPOSAL_ID,))
            self.REVIEWERS_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "SCIENCE_PARTY_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'LEAP_SP'", (self.PROPOSAL_ID,))
            self.SCIENCE_PARTY_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "REFERENCES_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'REFERENCES'", (self.PROPOSAL_ID,))
            self.REFERENCES_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "ENGAGEMENT_PLAN_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'ENGAGEMENT'", (self.PROPOSAL_ID,))
            self.ENGAGEMENT_PLAN_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "MANAGEMENT_PLAN_FILENAME":
            self.cur.execute("SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'MANAGEMENT'", (self.PROPOSAL_ID,))
            self.MANAGEMENT_PLAN_FILENAME = get_safely(self.cur.fetchone(), 0)

        elif table == "COVERSHEET_PROPONENT_MAP":
            self.cur.execute("SELECT * FROM coversheet_proponent_map WHERE proposal_id = %s ORDER BY ordering", (self.PROPOSAL_ID,))
            self.COVERSHEET_PROPONENT_MAP_ROWS = self.cur.fetchall()
            self.COVERSHEET_PROPONENT_MAP_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE":
            self.cur.execute("SELECT * FROM site WHERE proposal_id=%s ORDER BY ordering", (self.PROPOSAL_ID,))
            self.SITE_ROWS = self.cur.fetchall()
            self.SITE_COLS = [desc[0] for desc in self.cur.description]

        elif table == "PROPONENT":
            self.get_db_rows("COVERSHEET_PROPONENT_MAP")
            # Retrieve data from the proponent table by getting a tuple indicating all proponents of this proposal
            proponent_ids = tuple(proponent[self.COVERSHEET_PROPONENT_MAP_COLS.index('proponent_id')] for proponent in self.COVERSHEET_PROPONENT_MAP_ROWS)

            # Get the list of proponents from the coversheet_proponent_map table. First check that there is at least one proponent
            # to get to prevent crashing from a bad query. If there are no proponents then adjust the query so that it returns
            # an empty query rather than a NoneType so that the function that actually creates the proponent list doesn't realize
            # anything is wrong. Repeat this same process for the Principle Lead (i.e. check if they exist and if not then handle
            # it).
            if proponent_ids:
                self.cur.execute(sql.SQL(
                    """SELECT * FROM proponent JOIN coversheet_proponent_map ON
                    proponent.id = coversheet_proponent_map.proponent_id WHERE proponent.id IN {} ORDER BY
                    coversheet_proponent_map.ordering""").format(sql.Literal(proponent_ids)))
                self.PROPONENT_ROWS = self.cur.fetchall()
                self.PROPONENT_COLS = [desc[0] for desc in self.cur.description]
            else:
                self.cur.execute("SELECT * FROM proponent LIMIT 0")
                self.PROPONENT_ROWS = self.cur.fetchall()
                self.PROPONENT_COLS = [desc[0] for desc in self.cur.description]

        elif table == "LEAD_PROPONENT":
            self.get_db_rows("COVERSHEET_PROPONENT_MAP")
            # Narrow the list of proposals to locate just the principle lead and query this person's proponent data. Technically
            # this could be done via a list access of the PROPONENT_ROWS[] above but repetitive code is easier to understand.
            principle_lead_id = tuple(
                proponent[self.COVERSHEET_PROPONENT_MAP_COLS.index('proponent_id')] for proponent in
                self.COVERSHEET_PROPONENT_MAP_ROWS if
                proponent[self.COVERSHEET_PROPONENT_MAP_COLS.index('role')] == 'Principal Lead' or proponent[
                    self.COVERSHEET_PROPONENT_MAP_COLS.index('role')] == 'Principal Lead and Data Lead')

            if principle_lead_id:
                self.cur.execute(sql.SQL("SELECT * FROM proponent WHERE id IN {}").format(sql.Literal(principle_lead_id)))
                self.LEAD_PROPONENT_ROW = self.cur.fetchone()
                self.LEAD_PROPONENT_COLS = [desc[0] for desc in self.cur.description]
            else:
                self.cur.execute("SELECT * FROM proponent LIMIT 0")
                self.LEAD_PROPONENT_ROW = self.cur.fetchone()
                self.LEAD_PROPONENT_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE_OPERATIONAL_INFO":
            self.get_db_rows("SITE")
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            if site_numbers:
                # Retrieve data from the site_operation_info table by specifying the list of active sites for an IN
                # specifier
                self.cur.execute(sql.SQL(
                    """SELECT * FROM site_operational_info JOIN site ON site_operational_info.site_id = site.id WHERE
                    site_id IN {} ORDER BY site.ordering""").format(
                    sql.Literal(site_numbers)))
                self.SITE_OPERATIONAL_INFO_ROWS = self.cur.fetchall()
                self.SITE_OPERATIONAL_INFO_COLS = [desc[0] for desc in self.cur.description]
            else:
                self.cur.execute("SELECT * FROM site_operational_info LIMIT 0")
                self.SITE_OPERATIONAL_INFO_ROWS = self.cur.fetchall()
                self.SITE_OPERATIONAL_INFO_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE_UPLOAD_FILENAMES":
            self.get_db_rows("SITE")
            self.SITE_UPLOAD_FILENAMES = []
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            for number in site_numbers:
                self.cur.execute(
                    """SELECT filename_out FROM pdf_uploads WHERE proposal_id = %s AND form_type = 'SSF' AND
                    site_id = %s""", (self.PROPOSAL_ID, number))
                self.SITE_UPLOAD_FILENAMES.append(get_safely(self.cur.fetchone(), 0))

        elif table == "SITE_MEASUREMENTS":
            self.get_db_rows("SITE")
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            if site_numbers:
                self.cur.execute(sql.SQL(
                    "SELECT * FROM site_measurements JOIN site ON site_measurements.site_id = site.id WHERE site_id IN {} ORDER BY site.ordering").format(
                    sql.Literal(site_numbers)))
                self.SITE_MEASUREMENTS_ROWS = self.cur.fetchall()
                self.SITE_MEASUREMENTS_COLS = [desc[0] for desc in self.cur.description]
            else:
                self.cur.execute("SELECT * FROM site_measurements LIMIT 0")
                self.SITE_MEASUREMENTS_ROWS = self.cur.fetchall()
                self.SITE_MEASUREMENTS_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE_LITHOS":
            self.get_db_rows("SITE")
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            # We collect the SITE_LITHOS rows differently because we need to ensure that there is exactly one subset of SITE_LITHOS
            # records (arbitrary length subset but at least one record) for each SITE row. If we just query regularly we won't know
            # which exact sites are missing SITE_LITHOS records (we only know how many sites total are missing SITE_LITHOS records).
            self.SITE_LITHOS_COLS = []
            self.SITE_LITHOS_ROWS = []
            for i in range(0, len(self.SITE_ROWS)):
                self.cur.execute(sql.SQL(
                    "SELECT * FROM site_lithos JOIN site ON site_lithos.site_id = site.id WHERE site_id = %s"),
                    (site_numbers[i],))
                if len(self.SITE_LITHOS_COLS) == 0:
                    self.SITE_LITHOS_COLS = [desc[0] for desc in self.cur.description]
                result = self.cur.fetchall()
                # Sort the lithology rows based on the "min-depth" first and resolve ties by then sorting by the "max_depth" column
                # Empty values should appear after 0 in the listing, so we replace treat null values as being very large numbers.
                # These large values only control the logic. They won't be displayed in the final document. Note: Null values
                # shouldn't be appearing anyway because this input should be enforced and check client-side but it's good to be
                # aware of this anyway.
                if len(result) != 0:
                    self.SITE_LITHOS_ROWS.extend(sorted(result, key=lambda x: (
                        docx_format_number(x[self.SITE_LITHOS_COLS.index('min_depth')], 99999999),
                        docx_format_number(x[self.SITE_LITHOS_COLS.index('max_depth')], 99999999))))
                else:
                    # Create dummy site lithos row if this site doesn't have any site lithos records attached to it. We
                    # do this to ensure that every site has at least 1 lithos record attached to it (even if the record
                    # has no actual data).
                    self.SITE_LITHOS_ROWS.append(tuple(['' for i in range(len(self.SITE_LITHOS_COLS))]))

        elif table == "SSO_USERS":
            self.get_db_rows("PROPOSAL")
            if self.PROPOSAL_ROW:
                self.cur.execute("SELECT * FROM sso_users WHERE username=%s",
                            (get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("user_id"), True),))
            else:
                self.cur.execute("SELECT * FROM sso_users LIMIT 0")
            self.SSO_USERS_ROW = self.cur.fetchone()
            self.SSO_USERS_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE_DATASET_INFO":
            self.get_db_rows("SITE")
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            if site_numbers:
                self.cur.execute(sql.SQL(
                    """SELECT * FROM site_dataset_info JOIN site ON site_dataset_info.site_id = site.id WHERE site_id
                    IN {} ORDER BY site.ordering""").format(sql.Literal(site_numbers)))
            else:
                self.cur.execute("SELECT * FROM site_dataset_info LIMIT 0")
            self.SITE_DATASET_INFO_ROWS = self.cur.fetchall()
            self.SITE_DATASET_INFO_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SITE_POLLUTION_SAFETY":
            self.get_db_rows("SITE")
            site_numbers = tuple(site[self.SITE_COLS.index('id')] for site in self.SITE_ROWS)
            if site_numbers:
                self.cur.execute(sql.SQL(
                    """SELECT * FROM site_pollution_safety_hazards JOIN site on site_pollution_safety_hazards.site_id =
                    site.id WHERE site_id IN {} ORDER BY site.ordering""").format(sql.Literal(site_numbers)))
            else:
                self.cur.execute("SELECT * FROM site_pollution_safety_hazards LIMIT 0")
            self.SITE_POLLUTION_SAFETY_ROWS = self.cur.fetchall()
            self.SITE_POLLUTION_SAFETY_COLS = [desc[0] for desc in self.cur.description]

        elif table == "SRR_CHECKLIST":
            self.cur.execute("SELECT * FROM srr_checklist WHERE proposal_id = %s", (self.PROPOSAL_ID,))
            self.SRR_CHECKLIST_ROW = self.cur.fetchone()
            self.SRR_CHECKLIST_COLS = [desc[0] for desc in self.cur.description]

        else:
            raise RuntimeError("An invalid query name was specified.")

        # Record that we've performed this query so that we don't do it again needlessly in the future.
        self.tables_queried.append(table)


    def get_general_site_info_tags(self):
        self.get_db_rows("SITE")
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("COVERSHEET")
        self.get_db_rows("SITE_OPERATIONAL_INFO")
        self.get_db_rows("SITE_MEASUREMENTS")
        # Return the template dictionary now that all database values have been accessed

        return {
            'proposal_title': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("title"), True),
            'proposal_submitted_date': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("submit_date"), True),
            'previous_drilling': [docx_format_string(site[self.SITE_COLS.index('previous_drilling')]) for site in self.SITE_ROWS],
            'site_objective': [docx_format_string(site[self.SITE_COLS.index('site_objective')]) for site in self.SITE_ROWS],
            'name': [site[self.SITE_COLS.index('name')] for site in self.SITE_ROWS],
            'area': [site[self.SITE_COLS.index('area')] for site in self.SITE_ROWS],
            'jurisdiction': [site[self.SITE_COLS.index('jurisdiction')] for site in self.SITE_ROWS],
            'dist_to_land': [site[self.SITE_COLS.index('dist_to_land')] for site in self.SITE_ROWS],
            'former_sitename': [site[self.SITE_COLS.index('former_sitename')] for site in self.SITE_ROWS],
            'latitude': [docx_format_number(site[self.SITE_COLS.index('latitude')], val_to_return_if_null='') for site in self.SITE_ROWS],
            'longitude': [docx_format_number(site[self.SITE_COLS.index('longitude')], val_to_return_if_null='') for site in self.SITE_ROWS],
            'lat_long': [
                docx_format_number(site[self.SITE_COLS.index('latitude')], return_str=True, val_to_return_if_null='') + "\n" +
                docx_format_number(site[self.SITE_COLS.index('longitude')], return_str=True, val_to_return_if_null='') for site in self.SITE_ROWS],
            'datum': [site[self.SITE_COLS.index('datum')] for site in self.SITE_ROWS],
            'water_depth': [docx_format_number(site[self.SITE_COLS.index('water_depth')]) for site in self.SITE_ROWS],
            'is_primary_cb': [site[self.SITE_COLS.index('is_primary')] == 'primary' for site in self.SITE_ROWS],
            'is_alternate_cb': [site[self.SITE_COLS.index('is_primary')] == 'alternate' for site in self.SITE_ROWS],
            'site_operational_info.total_days_on_site': [docx_format_number(
                docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('days_drilling')], val_to_return_if_null=0, return_str=False) +
                docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('days_logging')], val_to_return_if_null=0, return_str=False),
                val_to_return_if_null=0) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_shallow_gas_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_shallow_gas')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_seabed_compl_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_seabed_compl')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_shall_water_flow_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_shall_water_flow')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_slide_turb_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_slide_turb')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_hydrotherm_act_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_hydrotherm_act')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_seabed_soft_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_seabed_soft')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_hc_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_hc')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_currents_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_currents')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_ch4h2o_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_ch4h2o')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_abnorm_p_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_abnorm_p')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_fract_zone_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_fract_zone')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_dia_volc_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_dia_volc')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_mm_object_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_mm_object')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_fault_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_fault')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_high_temp_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_high_temp')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_h2s_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_h2s')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_high_dip_angle_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_high_dip_angle')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_ice_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_ice')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_co2_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_co2')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.weather_win': [docx_format_string(site[self.SITE_OPERATIONAL_INFO_COLS.index('weather_win')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_sens_mar_habitat': [docx_format_string(site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_sens_mar_habitat')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.hw_other': [docx_format_string(site[self.SITE_OPERATIONAL_INFO_COLS.index('hw_other')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.days_drilling': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('days_drilling')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.days_logging': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('days_logging')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_fut': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_fut')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_apc_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_apc')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_xcb_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_xcb')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_rcb_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_rcb')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_reentry_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_reentry')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_pcs_cb': [site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_pcs')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.sediment_penetration': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('sediment_penetration')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.basement_penetration': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('basement_penetration')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.total_sediment': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('total_sediment')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.total_penetration': [docx_format_number(site[self.SITE_OPERATIONAL_INFO_COLS.index('total_penetration')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.plan_core': [docx_format_string(site[self.SITE_OPERATIONAL_INFO_COLS.index('plan_core')]) for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_measurement.wl_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('wl_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.mag_susc_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('mag_susc_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.form_img_ac_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('form_img_ac_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.bh_t_p_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('bh_t_p_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.vsp_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('vsp_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.lwd_acc_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('dens_neut_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.neut_poros_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('neut_poros_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.lith_dens_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('lith_dens_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.g_ray_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('g_ray_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.std_acc_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('std_acc_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.form_img_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('form_img_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.check_shot_survey_rel': [site[self.SITE_MEASUREMENTS_COLS.index('check_shot_survey_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.form_t_p_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('form_t_p_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.other_rel': [site[self.SITE_MEASUREMENTS_COLS.index('other')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.other_obj': [site[self.SITE_MEASUREMENTS_COLS.index('other_obj')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_measurement.resist_rel_cb': [site[self.SITE_MEASUREMENTS_COLS.index('resist_rel')] for site in self.SITE_MEASUREMENTS_ROWS],
            'site_operational_info.sediment_litho': [site[self.SITE_OPERATIONAL_INFO_COLS.index('sediment_litho')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
            'site_operational_info.basement_litho': [site[self.SITE_OPERATIONAL_INFO_COLS.index('basement_litho')] for site in self.SITE_OPERATIONAL_INFO_ROWS],
        }


    def get_site_survey_tags(self):
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("SITE_DATASET_INFO")
        return {
            'proposal_number': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_number"), True),
            'proposal_type_name': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_type"), True),
            'proposal_version': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("vers"), True),
            'site_dataset_info.primary_hrsr_will_upload': [site[self.SITE_DATASET_INFO_COLS.index('primary_hrsr_will_upload')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.crossing_hrsr_will_upload': [site[self.SITE_DATASET_INFO_COLS.index('crossing_hrsr_will_upload')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.primary_dpsr_will_upload': [site[self.SITE_DATASET_INFO_COLS.index('primary_dpsr_will_upload')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.crossing_dpsr_will_upload': [site[self.SITE_DATASET_INFO_COLS.index('crossing_dpsr_will_upload')] for site in self.SITE_DATASET_INFO_ROWS],

            'site_dataset_info.primary_hrsr_loc_pos_descp': [
                docx_format_seismic_reflection_data(
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_hrsr_location')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_hrsr_position')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_hrsr_position_type')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_hrsr_description')]
                ) for i in range(0, len(self.SITE_DATASET_INFO_ROWS))],

            'site_dataset_info.crossing_hrsr_loc_pos_descp': [
                docx_format_seismic_reflection_data(
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_hrsr_location')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_hrsr_position')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_hrsr_position_type')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_hrsr_description')]
                ) for i in range(0, len(self.SITE_DATASET_INFO_ROWS))],

            'site_dataset_info.primary_dpsr_loc_pos_descp': [
                docx_format_seismic_reflection_data(
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_dpsr_location')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_dpsr_position')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_dpsr_position_type')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('primary_dpsr_description')]
                ) for i in range(0, len(self.SITE_DATASET_INFO_ROWS))],

            'site_dataset_info.crossing_dpsr_loc_pos_descp': [
                docx_format_seismic_reflection_data(
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_dpsr_location')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_dpsr_position')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_dpsr_position_type')],
                    self.SITE_DATASET_INFO_ROWS[i][self.SITE_DATASET_INFO_COLS.index('crossing_dpsr_description')]
                ) for i in range(0, len(self.SITE_DATASET_INFO_ROWS))],
            'site_dataset_info.seism_veloc_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('seism_veloc_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.seism_veloc_dsc': [site[self.SITE_DATASET_INFO_COLS.index('seism_veloc_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.seismic_grid_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('seismic_grid_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.seismic_grid_dsc': [site[self.SITE_DATASET_INFO_COLS.index('seismic_grid_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.refraction_surf_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('refraction_surf_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.refraction_surf_dsc': [site[self.SITE_DATASET_INFO_COLS.index('refraction_surf_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.refraction_bottom_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('refraction_bottom_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.refraction_bottom_dsc': [site[self.SITE_DATASET_INFO_COLS.index('refraction_bottom_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.a_3_5_khz_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('a_3_5_khz_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.a_3_5_khz_dsc': [site[self.SITE_DATASET_INFO_COLS.index('a_3_5_khz_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.swath_bathy_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('swath_bathy_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.swath_bathy_dsc': [site[self.SITE_DATASET_INFO_COLS.index('swath_bathy_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.side_look_sonar_surf_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('side_look_sonar_surf_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.side_look_sonar_surf_dsc': [site[self.SITE_DATASET_INFO_COLS.index('side_look_sonar_surf_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.side_look_sonar_bottom_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('side_look_sonar_bottom_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.side_look_sonar_bottom_dsc': [site[self.SITE_DATASET_INFO_COLS.index('side_look_sonar_bottom_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.photo_video_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('photo_video_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.photo_video_dsc': [site[self.SITE_DATASET_INFO_COLS.index('photo_video_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.heat_flow_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('heat_flow_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.heat_flow_dsc': [site[self.SITE_DATASET_INFO_COLS.index('heat_flow_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.magnetics_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('magnetics_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.magnetics_dsc': [site[self.SITE_DATASET_INFO_COLS.index('magnetics_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.gravity_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('gravity_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.gravity_dsc': [site[self.SITE_DATASET_INFO_COLS.index('gravity_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.sedim_cores_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('sedim_cores_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.sedim_cores_dsc': [site[self.SITE_DATASET_INFO_COLS.index('sedim_cores_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.rock_samples_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('rock_samples_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.rock_samples_dsc': [site[self.SITE_DATASET_INFO_COLS.index('rock_samples_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.water_current_data_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('water_current_data_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.water_current_data_dsc': [site[self.SITE_DATASET_INFO_COLS.index('water_current_data_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.ice_cond_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('ice_cond_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.ice_cond_dsc': [site[self.SITE_DATASET_INFO_COLS.index('ice_cond_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.obs_micros_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('obs_micros_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.obs_micros_dsc': [site[self.SITE_DATASET_INFO_COLS.index('obs_micros_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.navigation_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('navigation_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.navigation_dsc': [site[self.SITE_DATASET_INFO_COLS.index('navigation_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.other_in_ssdb': [site[self.SITE_DATASET_INFO_COLS.index('other_in_ssdb')] for site in self.SITE_DATASET_INFO_ROWS],
            'site_dataset_info.other_dsc': [site[self.SITE_DATASET_INFO_COLS.index('other_dsc')] for site in self.SITE_DATASET_INFO_ROWS],
        }


    def get_environmental_protection_tags(self):
        self.get_db_rows("SITE_POLLUTION_SAFETY")
        return {
            'site_pollution_safety_hazard.oper_summary': [site[self.SITE_POLLUTION_SAFETY_COLS.index('oper_summary')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.hc_dsdp_odp': [site[self.SITE_POLLUTION_SAFETY_COLS.index('hc_dsdp_odp')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.hc_com': [site[self.SITE_POLLUTION_SAFETY_COLS.index('hc_com')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.ch4h2o': [site[self.SITE_POLLUTION_SAFETY_COLS.index('ch4h2o')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.hc_accum': [site[self.SITE_POLLUTION_SAFETY_COLS.index('hc_accum')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.precaut_sp': [site[self.SITE_POLLUTION_SAFETY_COLS.index('precaut_sp')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.aband_proc': [site[self.SITE_POLLUTION_SAFETY_COLS.index('aband_proc')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.mm_hazards': [site[self.SITE_POLLUTION_SAFETY_COLS.index('mm_hazards')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
            'site_pollution_safety_hazard.major_risk': [site[self.SITE_POLLUTION_SAFETY_COLS.index('major_risk')] for site in self.SITE_POLLUTION_SAFETY_ROWS],
        }


    def get_coversheet_tags(self):
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("COVERSHEET")
        self.get_db_rows("PROPONENT")
        self.get_db_rows("LEAD_PROPONENT")
        return {
            'generated_date': self.footer_text,
            'proposal_number': docx_format_string(get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_number"), True)),
            'proposal_type_name': docx_format_string(get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_type"), True)),
            'proposal_version': docx_format_string(get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("vers"), True)),
            'proponent_names': ', '.join(map(str, [self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('first')] + ' ' + self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('last')] for i in range(0, len(self.PROPONENT_ROWS))])),
            'principal_lead_name': get_safely(self.LEAD_PROPONENT_ROW, self.LEAD_PROPONENT_COLS.index("first"), True) + " " + get_safely(self.LEAD_PROPONENT_ROW, self.LEAD_PROPONENT_COLS.index("last"), True),
            'principal_lead_affiliation': get_safely(self.LEAD_PROPONENT_ROW, self.LEAD_PROPONENT_COLS.index("affiliation"), True),
            'principal_lead_country': get_safely(self.LEAD_PROPONENT_ROW, self.LEAD_PROPONENT_COLS.index("country"), True),
            'perm_to_post_cb': (get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("perm_to_post"), True), {'1': 'Yes', '0': 'No'}),
            'coversheet.title_short': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("title_short"), True),
            'coversheet.title': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("title"), True),
            'coversheet.keywords': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("keywords"), True),
            'coversheet.geo_area': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("geo_area"), True),
            'coversheet.abstract': docx_format_string(get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("abstract"), True)),
            'coversheet.objective': docx_format_string(get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("objective"), True)),
            'coversheet.contact_operator_ans': (get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("contact_operator_ans"), True), {'True': 'Yes', 'False': 'No', 'yes': 'Yes', 'no': 'No'}),
            'coversheet.non_stnd_measures': docx_format_string(get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("non_stnd_measures"), True)),
            'coversheet.sci_plain_lang': docx_format_string(get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("sci_plain_lang"), True)),
            'received_for_date': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("received_for_date"), True),
            'resubmission_prpsl_num': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("resubmission_prpsl_num"), True),
            'resubmission_explanation': docx_format_string(get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("resubmission_explanation"), True)),
        }


    def get_proposed_sites_tags(self):
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("SITE")
        return {
            'proposal_number': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_number"), True),
            'proposal_type_name': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_type"), True),
            'proposal_version': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("vers"), True),
            'total_sites': len(self.SITE_ROWS),
            'total_primary_sites': len(
                [site for site in self.SITE_ROWS if "primary" in site[self.SITE_COLS.index('is_primary')]]),
            'total_alt_sites': len([site for site in self.SITE_ROWS if "alternate" in site[self.SITE_COLS.index('is_primary')]]),
            'total_ns_sites': len([site for site in self.SITE_ROWS if "not set" in site[self.SITE_COLS.index('is_primary')]]),
        }


    def get_proponents_list_tags(self):
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("SSO_USERS")
        return {
            'proposal_number': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("proposal_number"), True),
            'contact_person_full_name': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("firstname"), True) + " " + get_safely(
                self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("lastname"), True),
            'contact_person_department': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("department"), True),
            'contact_person_org': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("organization"), True),
            'contact_person_addr': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("address"), True),
            'contact_person_city': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("city"), True),
            'contact_person_state': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("state"), True),
            'contact_person_zipcode': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("zipcode"), True),
            'contact_person_country': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("country"), True),
            'contact_person_email': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("email"), True),
            'contact_person_phone': get_safely(self.SSO_USERS_ROW, self.SSO_USERS_COLS.index("phone"), True),
        }


    def get_conditional_template_tags(self):
        self.get_db_rows("PROPOSAL")
        self.get_db_rows("COVERSHEET")
        return {
            'new_proposal': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("is_resubmission"), False) == 'true' or get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("resubmission_type"), True) != "",
            'coversheet.contact_operator_ans': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("contact_operator_ans"),False) is not None,
            'coversheet.sci_plain_lang': get_safely(self.COVERSHEET_ROW, self.COVERSHEET_COLS.index("sci_plain_lang"), False) is not None,
            'resubmission_type': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("resubmission_type"),False) == 'from_older_submission',
            'resubmission_prpsl_num': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("is_resubmission"), False) == 'true',
            'resubmission_from_declined_proposal': get_safely(self.PROPOSAL_ROW, self.PROPOSAL_COLS.index("resubmission_type"),False) == 'from_declined'
        }


    def generate_site_forms_full(self):
        """
        This function generates the full set of site forms including site forms 1, 2, 4, and 5. Each Word document
        is stored in the directory specified by PROPOSAL_DIR and they are ordered by their names which are 1.docx,
        2.docx, 3.docx, etc...
        :return:
        """
        filename0 = "iodp_proposal_pdf_site_info_template0.docx"
        filename1 = "iodp_proposal_pdf_site_info_template1.docx"
        filename2 = "iodp_proposal_pdf_site_info_template2.docx"
        filename3 = "iodp_proposal_pdf_site_survey_template.docx"
        filename4 = "iodp_proposal_pdf_site_env_protection_template.docx"
        filename5 = "iodp_proposal_pdf_site_lithologies_template.docx"

        # We want to map "True" values to a checkmark emoji so that the checkboxes are displayed with an emoji if true.
        # The hex digits are a utf-8 encoding of a black checkmark emoji.
        site_table_tf_display_dict = {'True': "☑", 'False': "☐", 'None': ''}
        site_survey_tf_display_dict = {'yes': 'yes️', 'no': 'no'}  # Update these values to automatically replace "yes"
        site_environ_display_dict = {"None": ""}  # Current proposals have a blank field rather than the text "None" so
        # we replace all "None" with the empty string.

        site_form_tags = {}
        for key, value in self.get_general_site_info_tags().items():
            site_form_tags[key] = (value, site_table_tf_display_dict)

        site_survey_tags = {}
        for key, value in self.get_site_survey_tags().items():
            if isinstance(value, list):  # Format all paragraphs in the "Details of available data..." section
                value = [docx_format_string(val) for val in value]
            site_survey_tags[key] = (value, site_survey_tf_display_dict)

        site_environ_tags = {}
        for key, value in self.get_environmental_protection_tags().items():
            site_environ_tags[key] = (value, site_environ_display_dict)

        joined_dicts = {**site_form_tags, **site_survey_tags, **site_environ_tags}

        # Get the SITE_LITHOS_ROWS member variable defined so we can read its data.
        self.get_db_rows("SITE_LITHOS")
        # Determine which entries of the lithos site table correspond to which site id by creating a list with
        # indices marking which section of SITE_LITHOS_ROWS corresponds to the first site, which to the second,
        # and so on...
        site_lithos_sections = [row[1] for row in self.SITE_LITHOS_ROWS]  # The lithologies are listed in order in accordance
        # with the site order so indicate where each lithologies section begins and ends for each site of this proposal.
        # Note that the site lithos records will be already sorted. This was done in the same block that queried them.
        seen_elements = set()  # This data structure tracks the which sites have had collected all of their lithos records.
        indices = []  # Indices define the length of the block for a site. Each value indicates the first index of the
        # next block of lithologies data for that position's corresponding site.
        if site_lithos_sections:
            for i, element in enumerate(site_lithos_sections):
                if element == "":
                    indices.append(i)  # If at this index the site lithos record is empty that means there were none for
                    # this site and that a dummy record was inserted so we mark the location of the dummy record.
                elif element not in seen_elements:
                    seen_elements.add(element)
                    indices.append(i)
            indices.append(len(site_lithos_sections))

        for i in range(0, len(self.SITE_ROWS)):
            counter = i * 6
            doc = Document(join(self.PDF_GEN_DIR, filename0))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 0) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 0) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename1))
            docx_define_styles(doc)
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 1) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 1) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename2))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 2) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 2) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename3))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 3) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 3) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename4))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 4) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 4) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename5))
            docx_define_styles(doc)
            docx_search_and_replace_tags(doc, joined_dicts, i)
            lithologies_data = []

            for j in range(indices[i], indices[i + 1]):  # For every site in a that's associated with this proposal
                site_lithos = ()
                if self.SITE_LITHOS_ROWS[j][1] != '':
                    site_lithos += (docx_format_number(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('min_depth')], "",
                                                       True)) + ' - ' + docx_format_number(
                        self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('max_depth')], "", True),
                else:
                    site_lithos += ('N/A',)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('key_event')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('age')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('velocity')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('lithology')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('paleo_env')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('sed_accum')]),)
                site_lithos += (docx_format_string(self.SITE_LITHOS_ROWS[j][self.SITE_LITHOS_COLS.index('comment')]),)
                lithologies_data.append(site_lithos)  # Add this site to the list of sites
            docx_build_table(doc, 0, lithologies_data)
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 5) + '.docx'))
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 5) + '.docx'))


    def generate_coversheet_page_full(self):
        COVERSHEET_TAGS = self.get_coversheet_tags()
        CONDITIONAL_TEMPLATE_TAGS = self.get_conditional_template_tags()

        doc0 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template0.docx"))
        doc1 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template1.docx"))
        doc2 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template2.docx"))
        doc3 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template3.docx"))
        doc4 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template4.docx"))

        docx_search_and_replace_tags(doc0, COVERSHEET_TAGS)
        docx_redact_conditional(doc0, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc1, COVERSHEET_TAGS)
        docx_redact_conditional(doc1, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc2, COVERSHEET_TAGS)
        docx_redact_conditional(doc2, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc3, COVERSHEET_TAGS)
        docx_redact_conditional(doc3, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc4, COVERSHEET_TAGS)
        docx_redact_conditional(doc4, CONDITIONAL_TEMPLATE_TAGS)

        doc0.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template0.docx"))
        doc1.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"))
        doc2.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2.docx"))
        doc3.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template3.docx"))
        doc4.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template4.docx"))


    def generate_coversheet_page_pre_leap(self):
        COVERSHEET_TAGS = self.get_coversheet_tags()
        CONDITIONAL_TEMPLATE_TAGS = self.get_conditional_template_tags()

        doc0 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template0.docx"))
        doc1 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template1.docx"))
        doc2 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template2_leap.docx"))

        docx_search_and_replace_tags(doc0, COVERSHEET_TAGS)
        docx_redact_conditional(doc0, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc1, COVERSHEET_TAGS)
        docx_redact_conditional(doc1, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc2, COVERSHEET_TAGS)
        docx_redact_conditional(doc2, CONDITIONAL_TEMPLATE_TAGS)

        doc0.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template0.docx"))
        doc1.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"))
        doc2.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"))


    def generate_coversheet_page_full_leap(self):
        COVERSHEET_TAGS = self.get_coversheet_tags()
        CONDITIONAL_TEMPLATE_TAGS = self.get_conditional_template_tags()

        doc0 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template0.docx"))
        doc1 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template1.docx"))
        doc2 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template2_leap.docx"))
        doc3 = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_coversheet_template4_leap.docx"))

        docx_search_and_replace_tags(doc0, COVERSHEET_TAGS)
        docx_redact_conditional(doc0, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc1, COVERSHEET_TAGS)
        docx_redact_conditional(doc1, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc2, COVERSHEET_TAGS)
        docx_redact_conditional(doc2, CONDITIONAL_TEMPLATE_TAGS)
        docx_search_and_replace_tags(doc3, COVERSHEET_TAGS)
        docx_redact_conditional(doc3, CONDITIONAL_TEMPLATE_TAGS)

        doc0.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template0.docx"))
        doc1.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"))
        doc2.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"))
        doc3.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template4_leap.docx"))



    def generate_proponents_page(self):
        self.get_db_rows("PROPONENT")
        self.get_db_rows("COVERSHEET_PROPONENT_MAP")
        doc = Document(join(self.PDF_GEN_DIR, "iodp_proposal_pdf_proponent_list_template.docx"))
        PROPONENTS_LIST_TAGS = self.get_proponents_list_tags()
        docx_define_styles(doc)
        docx_search_and_replace_tags(doc, PROPONENTS_LIST_TAGS)
        proponents_data = []
        if self.PROPONENT_ROWS:
            for i in range(0, len(self.PROPONENT_ROWS)):  # For every site in a that's associated with this proposal
                proponent = ()
                proponent += (docx_format_string(self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('first')]),)
                proponent += (docx_format_string(self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('last')]),)
                proponent += (docx_format_string(self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('affiliation')]),)
                proponent += (docx_format_string(self.PROPONENT_ROWS[i][self.PROPONENT_COLS.index('country')]),)
                proponent += (docx_format_string(self.COVERSHEET_PROPONENT_MAP_ROWS[i][self.COVERSHEET_PROPONENT_MAP_COLS.index('role')]),)
                proponent += (docx_format_string(self.COVERSHEET_PROPONENT_MAP_ROWS[i][self.COVERSHEET_PROPONENT_MAP_COLS.index('expertise')]),)
                proponents_data.append(proponent)  # Add this site to the list of sites
        else:
            proponents_data = [("N/A", "", "", "", "", "")]
        docx_build_table(doc, 1, proponents_data)
        doc.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_proponent_list_template.docx"))


    def generate_srr_checklist_page(self):
        self.get_db_rows("SRR_CHECKLIST")
        return {
            "srr_checklist.q1": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q1")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q2": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q2")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q3": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q3")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q4": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q4")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q5": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q5")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q6": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q6")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q7": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q7")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q8": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q8")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q9": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q9")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q10": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q10")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q11": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q11")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q12": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q12")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q13": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q13")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q14": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q14")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q15": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q15")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q16": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q16")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q17": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q17")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q18": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q18")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q19": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q19")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.q20": (
                get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("q20")),
                {"0": "No", "1": "Yes"}
            ),
            "srr_checklist.comments": get_safely(self.SRR_CHECKLIST_ROW, self.SRR_CHECKLIST_COLS.index("comments")),
        }


    def generate_proposed_sites_page(self):
        self.get_db_rows("SITE")
        self.get_db_rows("SITE_OPERATIONAL_INFO")
        doc = Document(self.PDF_GEN_DIR + "iodp_proposal_pdf_proposed_sites_template.docx")
        docx_define_styles(doc)
        docx_search_and_replace_tags(doc, self.get_proposed_sites_tags())
        proposed_sites_table_parsed = []
        if self.SITE_ROWS:
            for i in range(0, len(self.SITE_ROWS)):  # For every site in a that's associated with this proposal
                site = ()
                site += (get_safely(get_safely(self.SITE_ROWS, i), self.SITE_COLS.index('name')) + " " +
                     str(get_safely(get_safely(self.SITE_ROWS, i), self.SITE_COLS.index('is_primary'))),)
                site += (docx_format_number(get_safely(get_safely(self.SITE_ROWS, i), self.SITE_COLS.index('latitude')),
                    val_to_return_if_null='', return_str=True) + "\n" +
                    docx_format_number(get_safely(get_safely(self.SITE_ROWS, i), self.SITE_COLS.index('longitude')),
                    val_to_return_if_null='', return_str=True),)
                site += (
                    docx_format_number(get_safely(get_safely(self.SITE_ROWS, i), self.SITE_COLS.index('water_depth')),
                    val_to_return_if_null=0, return_str=True),)
                site += (
                    docx_format_number(get_safely(get_safely(self.SITE_OPERATIONAL_INFO_ROWS, i),
                    self.SITE_OPERATIONAL_INFO_COLS.index('sediment_penetration')),
                    val_to_return_if_null=0, return_str=True),)
                site += (
                    docx_format_number(get_safely(get_safely(self.SITE_OPERATIONAL_INFO_ROWS, i),
                    self.SITE_OPERATIONAL_INFO_COLS.index("basement_penetration")),
                    val_to_return_if_null=0, return_str=True),)
                site += (
                    docx_format_number(get_safely(get_safely(self.SITE_OPERATIONAL_INFO_ROWS, i),
                    self.SITE_OPERATIONAL_INFO_COLS.index('total_penetration')),
                    val_to_return_if_null=0, return_str=True),)
                site += (docx_format_string(get_safely(get_safely(self.SITE_ROWS, i),
                    self.SITE_COLS.index('site_objective'))),)
                proposed_sites_table_parsed.append(site)  # Add this site to the list of sites
        else:
            # Insert a dummy row if there are no sites to list.
            proposed_sites_table_parsed = [("N/A", "", "", "", "", "", "")]
        docx_build_table(doc, 0, proposed_sites_table_parsed)
        doc.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_proposed_sites_template.docx"))


    def get_page_identifiers(self):
        self.get_db_rows("COVERSHEET")
        header_text = self.get_coversheet_tags()['proposal_number'] + "-" + self.get_coversheet_tags()[
            'proposal_type_name'] + self.get_coversheet_tags()['proposal_version']
        # Because we are converting to PDF only once there's no point where we can inspect the number of pages of each
        # section individually. Thus the way we determine where the coversheet ends and proposed sites begins, for example,
        # is by inspecting text dumps for each page in output PDF and checking if those match specified sections from
        # each section. When we get a match then we know which section we are on in the for loop and bookmark the section
        # accordingly. The problem with this is that if we change things in the templates docx files we will need to be
        # sure to update the identifier variables here in the case that they change and no longer appear in the PDF.
        return {
            "header_text_identifier": header_text,
            "coversheet_page_identifier": header_text + "\nIODP Proposal Coversheet",
            "empty_page_identifier": header_text + "\n" + self.footer_text,
            "proposed_sites_page_identifier": header_text + "\nProposed Sites\n(Total proposed sites:",
            "proposed_sites_page_continued_identifier": header_text + "\nSite NamePosition\n(Lat, Lon)Water",
            "proponents_page_continued_identifier": header_text + ("\nFirst NameLast Name Affiliation CountryRoleExpertise\n"),
            "proponents_page_identifier": header_text + "\nContact Information\nContact Person:",
            "site_info_page_identifier": header_text + ("\nIODP Site Forms\nGeneral Site Information\nSection A: Proposal Information"),
            "site_survey_detail_identifier": header_text + "\nIODP Site Forms\nSite Survey Detail",
            "site_env_protection_identifier": header_text + "\nIODP Site Forms\nEnvironmental Protection",
            "site_lithologies_identifier": header_text + "\nIODP Site Forms\nLithologies",
            "safety_review_page_identifier": header_text + "\nSafety Review Preparation\nSeismic Data Presentation Question",
        }


    def remove_temp_files(self, exclude_list=[]):
        """
        This function deletes all files that were created in the PDF generation process and are intermediary files
        that are not part of the final product.
        :param exclude_list: This is a list of full file paths that should not be deleted even if their naming
        matches the pattern of a temporary file. This is useful if you are allowing the user to specify the output
        of the final PDF.
        :return: None
        """
        temp_files = glob.glob(join(self.PROPOSAL_DIR, "TEMP*"))
        site_files = glob.glob(join(self.PROPOSAL_DIR, "SITE*"))
        pdf_uploads_file_repairs = glob.glob(join(join(self.PROPOSAL_DIR, "pdf_uploads"), "TEMP*"))

        for file_path in temp_files:
            if file_path not in exclude_list:
                os.remove(file_path)
        for file_path in site_files:
            if file_path not in exclude_list:
                os.remove(file_path)
        for file_path in pdf_uploads_file_repairs:
            if file_path not in exclude_list:
                os.remove(file_path)


    def generate_site_forms_pre(self):
        filename0 = "iodp_proposal_pdf_site_info_template0.docx"
        filename1 = "iodp_proposal_pdf_site_info_template1.docx"
        filename2 = "iodp_proposal_pdf_site_info_template2.docx"

        # We want to map "True" values to a checkmark emoji so that the checkboxes are displayed with an emoji if true.
        site_table_tf_display_dict = {'True': "☑", 'False': "☐", 'None': ""}

        site_form_tags = {}
        for key, value in self.get_general_site_info_tags().items():
            site_form_tags[key] = (value, site_table_tf_display_dict)

        joined_dicts = {**site_form_tags}

        self.get_db_rows("SITE")
        self.get_db_rows("PROPOSAL")
        for i in range(0, len(self.SITE_ROWS)):
            counter = i * 3
            doc = Document(join(self.PDF_GEN_DIR, filename0))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 0) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 0) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename1))
            docx_define_styles(doc)
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 1) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 1) + '.docx'))

            doc = Document(join(self.PDF_GEN_DIR, filename2))
            docx_search_and_replace_tags(doc, joined_dicts, i)
            doc.save(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 2) + '.docx'))
            self.site_file_names.append(join(self.PROPOSAL_DIR, 'SITE_' + str(counter + 2) + '.docx'))


    def generate_safety_review_prep_page(self):
        doc = Document(self.PDF_GEN_DIR + "iodp_proposal_pdf_safety_review_prep_template.docx")
        docx_search_and_replace_tags(doc, self.generate_srr_checklist_page())
        doc.save(join(self.PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_safety_review_prep_template.docx"))