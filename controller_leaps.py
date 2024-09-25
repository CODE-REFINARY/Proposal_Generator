import argparse
from pdf_gen_helper_functions import *
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Composer_Document
from decouple import config
from WordProposalGenerator import WordProposalGenerator

# Extract the proposal ID command line arg to determine which proposal to generate a PDF for.
parser = argparse.ArgumentParser(description="Generate a PDF given a proposal ID")
parser.add_argument("proposal_id", type=str, help="This str specifies the proposal id to generate.")
parser.add_argument("--coversheet-only", "-c", action="store_true", help="This bool tells the program to produce only the coversheet.")
parser.add_argument("--output-filename", "-o", action="store", help="This argument should be a string which will be the name of the output PDF.")

args = parser.parse_args()
PROPOSAL_ID = args.proposal_id

# All files are accessed via absolute paths specified using this.
PDF_GEN_DIR = validate_path(config("PDF_GEN_DIR"))

# This command is run to invoke the LibreOffice daemon process which converts word documents to PDFs
LIBRE_OFFICE_PYTHON_INSTALLATION_PATH = validate_path(config("LIBRE_OFFICE_PYTHON_INSTALLATION_PATH"), True)
PROPOSALS_BASE_DIR = validate_path(config("PROPOSALS_BASE_DIR"))
PROPOSAL_DIR = validate_path(join(PROPOSALS_BASE_DIR, PROPOSAL_ID))
PDF_UPLOADS_DIR = validate_path(join(PROPOSAL_DIR, "pdf_uploads"))

obj = WordProposalGenerator(PROPOSAL_ID)

obj.get_db_rows("PROPOSAL")
p_type = obj.PROPOSAL_ROW[obj.PROPOSAL_COLS.index("proposal_type")]   # This variable str stores the proposal p_type
c_only = args.coversheet_only
output_name = args.output_filename

###############################################################################
# ------ Generating PDFs from Word Template with WordProposalGenerator ------ #
###############################################################################

# --------- Generate PDFs from Word Templates --------- #
print("-GENERATING COVER PAGE-")
if p_type == "Pre-LEAP":
    obj.generate_coversheet_page_pre_leap()
else:
    obj.generate_coversheet_page_full_leap()

if not c_only:
    print("-GENERATING PROPONENTS PAGE-")
    obj.generate_proponents_page()


###################################
# ------ THE MERGING STAGE ------ #
###################################

# Merge all Word Documents into a single PDF
print("-MERGING WORD DOCUMENTS-")

# Get the first Word document which is the coversheet.
master = Composer_Document(join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template0.docx"))
composer_obj = Composer(master)

# --------- Specify the names of Word templates --------- #

# Get the rest of the coversheet files in the correct order
if c_only:
    if p_type == "Pre-LEAP":
        completed_templates = [
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"),
        ]
    else:
        completed_templates = [
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template4_leap.docx"),
        ]
else:
    if p_type == "Pre-LEAP":
        completed_templates = [
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_proponent_list_template.docx"),
        ]
    else:
        completed_templates = [
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template1.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template2_leap.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_coversheet_template4_leap.docx"),
            join(PROPOSAL_DIR, "TEMP_iodp_proposal_pdf_proponent_list_template.docx"),
        ]

# --------- Concatenate all Word templates together into a big Word document --------- #

# Concatenate the other Word template docs to the (now complete) coversheet
for doc in completed_templates:
    if Document(doc).paragraphs:
        master.add_page_break()
    composer_obj.append(Composer_Document(doc))


#########################################################
# ------ THE ACTUAL WORD -> PDF CONVERSION STAGE ------ #
#########################################################

# We have finished concatenating the Word docs together.
print("-CONVERTING WORD TO PDF-")
composer_obj.save(join(PROPOSAL_DIR, "TEMP_final.docx"))

# Run the external tool to perform the singular Word to PDF conversion. The result is a temp file that still need
# to be bookmarked and have user uploads merged in.
docx_convert_word_to_pdf(join(PROPOSAL_DIR, 'TEMP_final.docx'), join(PROPOSAL_DIR, 'TEMP_final.pdf'))


###############################################################
# ------ The User Upload Merging and Bookmarking Stage ------ #
###############################################################

print("-MERGING PDFS-")
with open(join(PROPOSAL_DIR, "TEMP_final.pdf"), 'rb') as infile:

    reader = PdfReader(infile)
    if not c_only:
        # Open all user uploads. If a user upload doesn't exist, rather than crash the program simply print that
        # it wasn't found and proceed.
        obj.get_db_rows("MAIN_TEXT_FILENAME")
        obj.get_db_rows("REFERENCES_FILENAME")
        obj.get_db_rows("CURRICULUM_VITAE_FILENAME")
        obj.get_db_rows("ENGAGEMENT_PLAN_FILENAME")
        obj.get_db_rows("MANAGEMENT_PLAN_FILENAME")

        if p_type == "Full-LEAP":
            obj.get_db_rows("SCIENCE_PARTY_FILENAME")


        # --------- Define the document readers for all user uploads that correspond to each proposal p_type --------- #

        # Open all user uploads. If a user upload doesn't exist, rather than crash the program simply print that
        # it wasn't found and proceed.
        user_upload_main_text_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.MAIN_TEXT_FILENAME, "Main Text")
        user_upload_references_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.REFERENCES_FILENAME, "References")
        user_upload_cv_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.CURRICULUM_VITAE_FILENAME, "Curriculum Vitae")
        user_upload_engagement_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.ENGAGEMENT_PLAN_FILENAME, "Engagement Plan")
        user_upload_management_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.MANAGEMENT_PLAN_FILENAME, "Management Plan")

        if p_type == "Full-LEAP":
            user_upload_science_party_reader = instantiate_pdf_reader(PDF_UPLOADS_DIR, obj.SCIENCE_PARTY_FILENAME, "Science Party")

    writer = PdfWriter()
    writer.page_mode = "/UseOutlines"
    cur_i = 0  # "cursor index"
    cur_o = 0  # "cursor offset"

    ids = obj.get_page_identifiers()

    # --------- Merge the PDF pages from the old PDF composed of Word and Bookmark these pages --------- #

    # Merge coversheet, proponents, and proposed sites sheets in.
    if c_only:
        writer.add_outline_item("Cover Sheet", 0)
        for i in range(0, len(reader.pages)):
            writer.add_page(reader.pages[i])

    else:
        cur_i, cur_o, _ = docx_bookmark_and_process_pages(reader, writer, "Cover Sheet", cur_i, cur_o,
                                                          ids["coversheet_page_identifier"],
                                                          ids["empty_page_identifier"], None)

        cur_i, cur_o, _ = docx_bookmark_and_process_pages(reader, writer, "Proponents", cur_i, cur_o,
                                                      ids["proponents_page_identifier"],
                                                      ids["empty_page_identifier"],
                                                      ids["proponents_page_continued_identifier"])

        # --------- Merge the User Uploads In --------- #
        cur_o += docx_append_pages(user_upload_main_text_reader, writer, "Main Text", cur_i + cur_o)[0]
        cur_o += docx_append_pages(user_upload_engagement_reader, writer, "Engagement Plan", cur_i + cur_o)[0]
        cur_o += docx_append_pages(user_upload_management_reader, writer, "Management Plan", cur_i + cur_o)[0]
        cur_o += docx_append_pages(user_upload_references_reader, writer, "References", cur_i + cur_o)[0]
        if p_type == "Full-LEAP":
            cur_o += docx_append_pages(user_upload_science_party_reader, writer, "Science Party", cur_i + cur_o)[0]
        cur_o += docx_append_pages(user_upload_cv_reader, writer, "Curricula Vitae", cur_i + cur_o)[0]


###########################################################
# ------ Write The Result As The Final Output PDF  ------ #
###########################################################
if output_name:
    writer.write(join(PROPOSAL_DIR, output_name))
else:
    writer.write(join(PROPOSAL_DIR, PROPOSAL_ID + ".pdf"))


###########################################################
# ------ Write The Result As The Final Output PDF  ------ #
###########################################################

print("-------------------------")
print("-PDF GENERATION FINISHED-")
print("-------------------------")

# Get a list of files that start with 'TEMP' in the target directory
print("-CLEANING UP-")
# Clean up the intermediate files. If the user specified a filename then tell the cleanup function not to delete
# this file since the user might have named it similar to an intermediate file.
if output_name:
    obj.remove_temp_files(exclude_list=[join(PROPOSAL_DIR, output_name)])
else:
    obj.remove_temp_files()
